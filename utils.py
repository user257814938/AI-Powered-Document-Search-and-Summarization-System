# Objectif — Implémenter le backend RAG : extraction → chunking → embeddings → index FAISS → recherche → résumé → évaluation → persistance

# -----------------------------------------------------------------------------
# 1. IMPORTS & STRUCTURE DE DONNÉES (Base du projet)
# -----------------------------------------------------------------------------

# Étape 1.1 — Imports techniques indispensables
from __future__ import annotations                                              # from : mot-clé d'import source | __future__ : module interne de compatibilité | import : mot-clé d'import cible | annotations : fonctionnalité de typage différé
from typing import List, Sequence                                               # from : mot-clé d'import source | typing : module de typage | import : mot-clé d'import cible | List : type liste générique | Sequence : type séquence générique
from dataclasses import dataclass                                               # from : mot-clé d'import source | dataclasses : module utilitaire pour classes | import : mot-clé d'import cible | dataclass : décorateur de classe

# Étape 1.2 — Structure de données centrale
@dataclass                                                                      # @dataclass : décorateur pour générer automatiquement les méthodes spéciales (__init__, __repr__)
class IndexedChunk:                                                             # class : mot-clé définition classe | IndexedChunk : nom de la classe définie
    text: str                                                                   # text : attribut pour le contenu | : : séparateur de type | str : type chaîne de caractères
    doc_id: str                                                                 # doc_id : attribut identifiant document | : : séparateur de type | str : type chaîne de caractères
    chunk_id: int                                                               # chunk_id : attribut identifiant segment | : : séparateur de type | int : type nombre entier


# -----------------------------------------------------------------------------
# 2. EXTRACTION DE DONNÉES (Entrée du pipeline)
# -----------------------------------------------------------------------------

# Étape 2.1 — Imports nécessaires pour lire les fichiers
from pathlib import Path                                                        # from : mot-clé d'import source | pathlib : module de gestion des chemins | import : mot-clé d'import cible | Path : classe objet chemin
from PyPDF2 import PdfReader                                                    # from : mot-clé d'import source | PyPDF2 : librairie de lecture PDF | import : mot-clé d'import cible | PdfReader : classe lecteur PDF
from docx import Document                                                       # from : mot-clé d'import source | docx : librairie de lecture Word | import : mot-clé d'import cible | Document : classe document Word

# Étape 2.2 — Logique d'extraction
def extract_text(file_path: Path) -> str:                                       # def : mot-clé fonction | extract_text : nom de la fonction | file_path : paramètre chemin | : : séparateur type | Path : type attendu | -> : flèche retour | str : type retourné
    """Extrait le texte brut depuis .txt, .pdf ou .docx."""
    suffix = file_path.suffix.lower()                                           # suffix : variable extension | = : opérateur d'affectation | file_path : objet chemin | .suffix : attribut extension | .lower() : méthode conversion minuscules
    if suffix == ".txt":                                                        # if : mot-clé condition | suffix : variable testée | == : opérateur égalité | ".txt" : chaîne littérale extension
        return file_path.read_text(encoding="utf-8", errors="ignore")           # return : mot-clé retour | file_path : objet chemin | .read_text : méthode lecture texte | encoding="utf-8" : argument encodage | errors="ignore" : argument gestion erreurs
    if suffix == ".pdf":                                                        # if : mot-clé condition | suffix : variable testée | == : opérateur égalité | ".pdf" : chaîne littérale extension
        reader = PdfReader(str(file_path))                                      # reader : variable lecteur | = : opérateur affectation | PdfReader : constructeur classe | str(file_path) : chemin converti en chaîne
        return "\n".join(page.extract_text() or "" for page in reader.pages)    # return : mot-clé retour | "\n" : saut de ligne | .join : méthode concaténation | page.extract_text() : extraction texte page | or "" : alternative si vide | for : boucle génératrice | page : variable itérateur | in : dans | reader.pages : liste des pages
    if suffix == ".docx":                                                       # if : mot-clé condition | suffix : variable testée | == : opérateur égalité | ".docx" : chaîne littérale extension
        doc = Document(str(file_path))                                          # doc : variable document | = : opérateur affectation | Document : constructeur classe | str(file_path) : chemin converti en chaîne
        return "\n".join(p.text for p in doc.paragraphs)                        # return : mot-clé retour | "\n" : saut de ligne | .join : méthode concaténation | p.text : attribut texte paragraphe | for : boucle génératrice | p : variable itérateur | in : dans | doc.paragraphs : liste des paragraphes
    raise ValueError(f"Extension non supportée: {suffix}")                      # raise : mot-clé lever exception | ValueError : type d'erreur | f"..." : chaîne formatée | {suffix} : injection variable suffixe


# -----------------------------------------------------------------------------
# 3. CHUNKING (Préparation du texte)
# -----------------------------------------------------------------------------

# Étape 3.1 — Import du Tokenizer
from transformers import AutoTokenizer                                          # from : mot-clé d'import source | transformers : librairie Hugging Face | import : mot-clé d'import cible | AutoTokenizer : classe de chargement automatique tokenizer

# Étape 3.2 — Logique de découpage
def build_tokenizer(model_name: str = "bert-base-uncased"):                     # def : mot-clé fonction | build_tokenizer : nom fonction | model_name : nom paramètre | : : séparateur type | str : type chaîne | = : valeur par défaut | "bert-base-uncased" : nom modèle par défaut
    return AutoTokenizer.from_pretrained(model_name)                            # return : mot-clé retour | AutoTokenizer : classe | .from_pretrained : méthode chargement modèle | model_name : argument nom modèle

def chunk_text(                                                                 # def : mot-clé définition fonction | chunk_text : nom de la fonction
    text: str,                                                                  # text : paramètre contenu texte | : : séparateur type | str : type chaîne de caractères
    tokenizer=None,                                                             # tokenizer : paramètre outil tokenisation | = : valeur par défaut | None : objet nul
    chunk_size: int = 250,                                                      # chunk_size : paramètre taille morceau | : : séparateur type | int : type entier | = : valeur par défaut | 250 : nombre de tokens
    overlap: int = 30                                                           # overlap : paramètre recouvrement | : : séparateur type | int : type entier | = : valeur par défaut | 30 : nombre de tokens
) -> List[str]:                                                                 # ) : fin paramètres | -> : flèche type retour | List : type liste | [str] : contenant des chaînes
    """Découpe le texte en segments (chunks) avec recouvrement."""
    tokenizer = tokenizer or build_tokenizer()                                  # tokenizer : variable locale | = : affectation | tokenizer : valeur actuelle | or : ou logique (si None) | build_tokenizer() : appel fonction construction
    tokens = tokenizer.encode(text, add_special_tokens=False)                   # tokens : liste identifiants numériques | = : affectation | tokenizer : objet outil | .encode : méthode conversion | text : entrée texte | add_special_tokens=False : argument sans balises spéciales
    chunks = []                                                                 # chunks : variable liste résultats | = : affectation | [] : initialisation liste vide
    step = max(chunk_size - overlap, 1)                                         # step : variable pas d'avancement | = : affectation | max : fonction maximum | chunk_size : taille | - : moins | overlap : recouvrement | , : séparateur arguments | 1 : valeur minimale

    for start in range(0, len(tokens), step):                                   # for : boucle itérative | start : variable index début | in : dans | range : générateur séquence | 0 : début | len(tokens) : fin (longueur totale) | step : pas d'incrément
        piece = tokens[start : start + chunk_size]                              # piece : segment courant tokens | = : affectation | tokens : liste source | [ : début slice | start : index début | : : séparateur slice | start + chunk_size : index fin | ] : fin slice
        if not piece:                                                           # if : condition | not : opérateur négation (si vide) | piece : variable testée
            continue                                                            # continue : mot-clé passer à l'itération suivante
        decoded = tokenizer.decode(piece, skip_special_tokens=True)             # decoded : texte reconstruit | = : affectation | tokenizer : objet outil | .decode : méthode inverse | piece : segment tokens | skip_special_tokens=True : argument ignorer balises
        if decoded.strip():                                                     # if : condition | decoded : variable texte | .strip() : méthode nettoyage espaces (test si non vide)
            chunks.append(decoded.strip())                                      # chunks : liste résultats | .append : méthode ajout élément | decoded : texte | .strip() : texte nettoyé

    return chunks                                                               # return : mot-clé retour fonction | chunks : liste finale des segments


# -----------------------------------------------------------------------------
# 4. EMBEDDINGS (Vectorisation)
# -----------------------------------------------------------------------------

# Étape 4.1 — Imports pour le calcul vectoriel
from sentence_transformers import SentenceTransformer                           # from : mot-clé d'import source | sentence_transformers : librairie SBERT | import : mot-clé d'import cible | SentenceTransformer : classe modèle embedding
import numpy as np                                                              # import : mot-clé chargement module | numpy : librairie calcul numérique | as : mot-clé alias | np : nom raccourci local

# Étape 4.2 — Logique d'encodage
def build_embedder(model_name: str = "sentence-transformers/all-MiniLM-L6-v2"): # def : mot-clé fonction | build_embedder : nom fonction | model_name : paramètre nom | : : type | str : chaîne | = : défaut | "..." : identifiant modèle
    return SentenceTransformer(model_name, device="cpu")                        # return : mot-clé retour | SentenceTransformer : constructeur classe | model_name : argument modèle | device="cpu" : argument forcer exécution processeur

def encode_chunks(                                                              # def : mot-clé définition fonction | encode_chunks : nom de la fonction
    chunks: Sequence[str],                                                      # chunks : paramètre entrée segments | : : type | Sequence : interface séquence | [str] : contenant chaînes
    embedder=None,                                                              # embedder : paramètre outil encodage | = : défaut | None : valeur nulle
    batch_size: int = 4                                                         # batch_size : paramètre taille lot | : : type | int : entier | = : défaut | 4 : nombre éléments
) -> np.ndarray:                                                                # ) : fin paramètres | -> : flèche retour | np : alias numpy | .ndarray : type tableau multidimensionnel
    """Convertit les textes en vecteurs numériques."""
    embedder = embedder or build_embedder()                                     # embedder : variable | = : affectation | embedder : valeur actuelle | or : fallback | build_embedder() : appel construction par défaut
    embeddings = embedder.encode(                                               # embeddings : variable résultat matrice | = : affectation | embedder : objet modèle | .encode : méthode calcul vecteurs
        list(chunks),                                                           # list : conversion en liste | chunks : variable entrée
        batch_size=batch_size,                                                  # batch_size : argument nommé | batch_size : valeur variable
        normalize_embeddings=True,                                              # normalize_embeddings : argument nommé normalisation | = : assignation | True : booléen vrai
        convert_to_numpy=True,                                                  # convert_to_numpy : argument nommé format sortie | = : assignation | True : booléen vrai
        show_progress_bar=False,                                                # show_progress_bar : argument nommé affichage | = : assignation | False : booléen faux
    )                                                                           # ) : fin appel fonction
    return embeddings.astype("float32")                                         # return : mot-clé retour | embeddings : matrice calculée | .astype : méthode conversion type | "float32" : format flottant 32 bits


# -----------------------------------------------------------------------------
# 5. INDEXATION & STOCKAGE (Base de données vectorielle)
# -----------------------------------------------------------------------------

# Étape 5.1 — Import du moteur de recherche vectoriel
import faiss                                                                    # import : mot-clé chargement module | faiss : librairie recherche similarité

# Étape 5.2 — Gestion de l'index
def build_faiss_index(embeddings: np.ndarray) -> faiss.IndexFlatL2:             # def : mot-clé fonction | build_faiss_index : nom | embeddings : paramètre vecteurs | : : type | np.ndarray : tableau numpy | -> : retour | faiss.IndexFlatL2 : type index spécifique
    """Crée un index FAISS pour la recherche de similarité L2."""
    if embeddings.ndim != 2:                                                    # if : condition | embeddings : tableau | .ndim : propriété nombre dimensions | != : opérateur différent | 2 : valeur attendue
        raise ValueError("Les embeddings doivent être de forme (n, dim)")       # raise : mot-clé lever erreur | ValueError : classe exception | "..." : message erreur
    dim = embeddings.shape[1]                                                   # dim : variable dimension vecteurs | = : affectation | embeddings : tableau | .shape : tuple dimensions | [1] : accès second élément (colonnes)
    index = faiss.IndexFlatL2(dim)                                              # index : variable objet index | = : affectation | faiss : module | .IndexFlatL2 : constructeur classe distance L2 | dim : argument dimension
    index.add(embeddings)                                                       # index : objet index | .add : méthode ajout données | embeddings : tableau vecteurs
    return index                                                                # return : mot-clé retour | index : objet index construit

def save_index(index: faiss.IndexFlatL2, path: Path) -> None:                   # def : fonction | save_index : nom | index : paramètre objet | path : paramètre chemin | -> : retour | None : vide
    path.parent.mkdir(parents=True, exist_ok=True)                              # path : objet chemin | .parent : attribut dossier parent | .mkdir : méthode création dossier | parents=True : création récursive | exist_ok=True : pas d'erreur si existe
    faiss.write_index(index, str(path))                                         # faiss : module | .write_index : fonction écriture disque | index : objet à sauvegarder | str(path) : chemin converti en chaîne

def load_index(path: Path) -> faiss.IndexFlatL2:                                # def : fonction | load_index : nom | path : paramètre chemin | -> : retour | faiss.IndexFlatL2 : type attendu
    return faiss.read_index(str(path))                                          # return : retour | faiss : module | .read_index : fonction lecture disque | str(path) : chemin converti en chaîne


# -----------------------------------------------------------------------------
# 6. RECHERCHE (Interrogation)
# -----------------------------------------------------------------------------

# Étape 6.1 — Logique de recherche (utilise FAISS importé au dessus)
def search(embeddings: np.ndarray, index: faiss.IndexFlatL2, top_k: int = 5):   # def : fonction | search : nom | embeddings : vecteur requête | index : base vectorielle | top_k : nombre résultats | : : type | int : entier | = : défaut | 5 : valeur
    """Trouve les k vecteurs les plus proches dans l'index."""
    scores, idxs = index.search(embeddings, top_k)                              # scores : variable distances | , : séparateur unpacking | idxs : variable indices | = : affectation | index : objet index | .search : méthode recherche | embeddings : requête | top_k : limite
    return scores, idxs                                                         # return : retour valeurs | scores : tableau distances | , : séparateur | idxs : tableau identifiants


# -----------------------------------------------------------------------------
# 7. RÉSUMÉ (Génération de réponse)
# -----------------------------------------------------------------------------

# Étape 7.1 — Import du pipeline de résumé
from transformers import pipeline                                               # from : mot-clé import source | transformers : librairie | import : mot-clé import cible | pipeline : fonction outil haut niveau

# Étape 7.2 — Logique de résumé
def build_summarizer(model_name: str = "t5-small"):                             # def : fonction | build_summarizer : nom | model_name : paramètre | : : type | str : chaîne | = : défaut | "t5-small" : modèle Google T5 léger
    return pipeline(                                                            # return : retour | pipeline : appel constructeur pipeline
        "summarization",                                                        # "summarization" : argument tâche NLP visée
        model=model_name,                                                       # model : argument nommé | = : assignation | model_name : variable
        device=-1,                                                              # device : argument matériel | = : assignation | -1 : valeur pour CPU
        truncation=True,                                                        # truncation : argument coupe texte | = : assignation | True : activé
    )                                                                           # ) : fin appel

def summarize_chunks(                                                           # def : définition fonction | summarize_chunks : nom
    chunks: Sequence[IndexedChunk],                                             # chunks : paramètre entrée | : : type | Sequence : suite générique | [IndexedChunk] : objets chunks
    summarizer=None,                                                            # summarizer : paramètre outil | = : défaut | None : vide
    max_length: int = 200,                                                      # max_length : paramètre longueur max | : : type | int : entier | = : défaut | 200 : tokens
    min_length: int = 30                                                        # min_length : paramètre longueur min | : : type | int : entier | = : défaut | 30 : tokens
) -> str:                                                                       # ) : fin params | -> : retour | str : chaîne texte
    """Génère un résumé à partir des chunks récupérés."""
    summarizer = summarizer or build_summarizer()                               # summarizer : variable | = : assignation | summarizer : existant | or : ou | build_summarizer() : nouveau
    merged = "\n".join(chunk.text for chunk in chunks)                          # merged : texte fusionné | = : assignation | "\n" : séparateur saut ligne | .join : méthode jointure | chunk.text : attribut texte | for : boucle génératrice | chunk : variable | in : dans | chunks : liste
    summary = summarizer(                                                       # summary : variable résultat brut | = : assignation | summarizer : appel fonction modèle
        merged,                                                                 # merged : argument texte entrée
        max_length=max_length,                                                  # max_length : argument contrainte max | = : valeur
        min_length=min_length,                                                  # min_length : argument contrainte min | = : valeur
        no_repeat_ngram_size=3,                                                 # no_repeat_ngram_size : argument anti-répétition | = : assignation | 3 : taille ngram interdits
        truncation=True,                                                        # truncation : argument coupe entrée | = : assignation | True : activé
    )                                                                           # ) : fin appel
    return summary[0]["summary_text"].strip()                                   # return : retour | summary : liste résultat | [0] : premier élément | ["summary_text"] : clé dictionnaire | .strip() : méthode nettoyage


# -----------------------------------------------------------------------------
# 8. ÉVALUATION & PERSISTANCE MÉTADONNÉES (Étapes finales)
# -----------------------------------------------------------------------------

# Étape 8.1 — Import pour la sauvegarde JSON
import json                                                                     # import : mot-clé chargement module | json : librairie format données

# Étape 8.2 — Métriques et sauvegarde
def precision_recall_at_k(                                                      # def : définition fonction | precision_recall_at_k : nom
    relevant_ids: Sequence[int],                                                # relevant_ids : paramètre identifiants pertinents | : : type | Sequence : suite | [int] : entiers
    retrieved_ids: Sequence[int],                                               # retrieved_ids : paramètre identifiants trouvés | : : type | Sequence : suite | [int] : entiers
    k: int = 5                                                                  # k : paramètre seuil coupure | : : type | int : entier | = : défaut | 5 : valeur
):                                                                              # ) : fin paramètres | : : début bloc
    relevant_set = set(relevant_ids)                                            # relevant_set : ensemble unique pertinents | = : affectation | set : constructeur ensemble | relevant_ids : liste entrée
    retrieved_top_k = retrieved_ids[:k]                                         # retrieved_top_k : sous-liste top k | = : affectation | retrieved_ids : liste complète | [:k] : slice premiers k éléments
    true_positive = len([rid for rid in retrieved_top_k if rid in relevant_set])# true_positive : nombre corrects | = : affectation | len : longueur | [ : début liste compréhension | rid : variable id | for : boucle | ... : filtre inclusion
    precision = true_positive / max(k, 1)                                       # precision : variable score précision | = : calcul | true_positive : numérateur | / : division | max : sécurité division zéro | k : dénominateur
    recall = true_positive / max(len(relevant_set), 1)                          # recall : variable score rappel | = : calcul | true_positive : numérateur | / : division | max : sécurité division zéro | len : total pertinents
    return precision, recall                                                    # return : retour valeurs | precision : score 1 | , : séparateur | recall : score 2

def save_metadata(chunks: Sequence[IndexedChunk], path: Path) -> None:          # def : fonction | save_metadata : nom | chunks : données entrée | path : chemin fichier | -> : retour | None : vide
    path.parent.mkdir(parents=True, exist_ok=True)                              # path.parent : dossier parent | .mkdir : création dossier | parents=True : récursif | exist_ok=True : pas d'erreur
    with path.open("w", encoding="utf-8") as f:                                 # with : gestionnaire contexte | path.open : ouverture fichier | "w" : mode écriture | encoding : utf8 | as : alias | f : variable fichier
        json.dump(                                                              # json : module | .dump : méthode sérialisation vers fichier
            [chunk.__dict__ for chunk in chunks],                               # [ : début liste | chunk.__dict__ : conversion objet en dictionnaire | for : boucle | chunk : variable | in : dans | chunks : source | ] : fin liste
            f,                                                                  # f : objet fichier destination
            ensure_ascii=False,                                                 # ensure_ascii : option encodage | = : assignation | False : autoriser caractères accentués
            indent=2,                                                           # indent : option formatage | = : assignation | 2 : nombre espaces indentation
        )                                                                       # ) : fin appel

def load_metadata(path: Path) -> List[IndexedChunk]:                            # def : fonction | load_metadata : nom | path : chemin fichier | -> : retour | List : type liste | [IndexedChunk] : objets
    data = json.loads(path.read_text(encoding="utf-8"))                         # data : variable données brutes | = : affectation | json.loads : méthode parsing chaîne | path.read_text : lecture contenu fichier
    return [IndexedChunk(**item) for item in data]                              # return : retour | [ : début liste | IndexedChunk : constructeur | **item : unpacking dictionnaire en arguments | for : boucle | item : variable | in : dans | data : source | ] : fin liste
